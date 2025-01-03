import docx
from django.http import FileResponse
from django.shortcuts import render
from utils import convert_docx_to_pdf


def main_index(request):
    """Главная страница"""
    context = {"title": "Главная"}
    return render(
        request,
        "base.html",
        {
            "context": context,
        },
    )


def get_pdf(request):
    """Отчёт в pdf"""
    print(request.POST.dict())
    data = request.POST.dict()
    del data["csrfmiddlewaretoken"]
    doc = docx.Document()
    doc.add_heading("Отчёт", 0)

    row_count = 0
    doc.add_heading("Исходные данные", 2)
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    for key, value in data.items():
        if key == "new_table":
            doc.add_heading(value, 2)
            table = doc.add_table(rows=15, cols=2)
            table.style = "Table Grid"
            row_count = 0
            continue
        cell = table.cell(row_count, 0)
        cell.text = str(key)
        cell = table.cell(row_count, 1)
        cell.text = str(value)
        row_count += 1

    doc.save("input.docx")
    convert_docx_to_pdf("input.docx", "output.pdf")

    return FileResponse(open("output.pdf", "rb"))
